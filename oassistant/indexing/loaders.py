
# oassistant/index/loaders.py

""" 
Load documents for later embedding, include metadata.
Supported 
    *.txt, *.md
    *.html
    *.pdf
    *.docx  note: convert *.doc to docx first
    *.odt   note: does not work well, convert to *.docx
    *.json  note: load conversations stored as json   
    todo:  should add xml 
"""

import json
from pathlib import Path

from langchain_core.documents import Document
from langchain_community.document_loaders import (
#    TextLoader,
#    UnstructuredODTLoader,
    PyPDFLoader,
)
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from markdown_it import MarkdownIt
import subprocess
import tempfile


def load_html_with_metadata(path, corpus_name):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        soup = BeautifulSoup(f, "lxml")

    title = (
        soup.title.get_text(strip=True)
        if soup.title and soup.title.get_text(strip=True)
        else path.stem
    )

    docs = []

    for header in soup.find_all(["h1", "h2", "h3"]):
        section_title = header.get_text(strip=True)

        content = []
        for sibling in header.next_siblings:
            if hasattr(sibling, "name") and sibling.name in ["h1", "h2", "h3"]:

                break
            if hasattr(sibling, "get_text"):
                content.append(sibling.get_text(" ", strip=True))

        text = "\n".join(content).strip()

        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": path.name,
                        "title": title,
                        "section": section_title,
                        "page": None,
                        "file_type": "html",
                        "corpus": corpus_name,
                    }
                )
            )

    return docs
    
# =============================
# Markdown(md) AST Loader
# =============================

def load_markdown_ast(path: Path, corpus_name):
    docs = []

    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    if not text.strip():
        return docs

    md = MarkdownIt()
    tokens = md.parse(text)

    title = path.stem
    current_headings = {1: None, 2: None, 3: None}
    buffer = []
    inside_code_block = False
    code_language = None

    def current_section():
        # Use deepest available heading
        for level in reversed([1, 2, 3]):
            if current_headings[level]:
                return current_headings[level]
        return "Introduction"

    def flush_section():
        if buffer:
            docs.append(
                Document(
                    page_content="\n".join(buffer).strip(),
                    metadata={
                        "source": path.name,
                        "file_type": "md",
                        "title": title,
                        "section": current_section(),
                        "page": None,
                        "contains_code": inside_code_block,
                        "code_language": code_language,
                        "corpus": corpus_name,
                    },
                )
            )

    i = 0
    while i < len(tokens):
        token = tokens[i]

        # ---------------- Headings ----------------
        if token.type == "heading_open":
            flush_section()
            buffer.clear()

            level = int(token.tag[1])  # h1, h2, h3
            heading_text = tokens[i + 1].content

            if level == 1:
                title = heading_text
                current_headings = {1: heading_text, 2: None, 3: None}
            else:
                current_headings[level] = heading_text

            i += 2  # skip inline + heading_close

        # ---------------- Paragraph / Inline ----------------
        elif token.type == "inline":
            buffer.append(token.content)

        # ---------------- Code Block ----------------
        elif token.type == "fence":
            flush_section()
            buffer.clear()

            inside_code_block = True
            code_language = token.info.strip() or None

            docs.append(
                Document(
                    page_content=token.content,
                    metadata={
                        "source": path.name,
                        "file_type": "md",
                        "title": title,
                        "section": current_section(),
                        "page": None,
                        "contains_code": True,
                        "code_language": code_language,
                        "corpus": corpus_name,
                    },
                )
            )

            inside_code_block = False
            code_language = None

        # ---------------- Lists ----------------
        elif token.type == "list_item_open":
            buffer.append("- ")

        i += 1

    flush_section()

    return docs   

# =============================
# Custom DOCX Loader (with sections + title)
# =============================

def load_docx_with_sections(path: Path, corpus_name):
    docs = []
    doc = DocxDocument(str(path))

    # Extract Word metadata title (fallback to filename)
    core_props = doc.core_properties
    title = core_props.title if core_props.title else path.stem

    current_section = "Introduction"

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style.name.startswith("Heading"):
            current_section = text
            continue

        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": path.name,
                    "file_type": "docx",
                    "title": title,
                    "section": current_section,
                    "corpus": corpus_name,
                },
            )
        )

    return docs

# =============================
# JSON Conversation Loader
# =============================

def load_json_conversation(path: Path, corpus_name):
    docs = []

    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    session_id = data.get("session_id", path.stem)
    messages = data.get("messages", [])

    i = 0
    while i < len(messages):
        msg = messages[i]

        # Pair user + assistant if possible
        if msg.get("role") == "user":
            user_content = msg.get("content", "")
            assistant_content = ""

            if i + 1 < len(messages) and messages[i + 1].get("role") == "assistant":
                assistant_content = messages[i + 1].get("content", "")
                i += 1  # skip assistant

            combined = (
                f"User:\n{user_content}\n\n"
                f"Assistant:\n{assistant_content}"
            )

            docs.append(
                Document(
                    page_content=combined.strip(),
                    metadata={
                        "source": path.name,
                        "file_type": "json_conversation",
                        "session_id": session_id,
                        "page": None,
                        "corpus": corpus_name,  # explicitly different ="memory" - in function call
                    },
                )
            )

        i += 1

    return docs
    

# =============================
# Step 1: Load documents
# =============================

def load_documents(folder: Path, corpus_name):
    docs = []

    for path in folder.glob("**/*"):
        suffix = path.suffix.lower()

        try:

            # ---------------- TXT ----------------
            if suffix == ".txt":

                with open(path, "r", encoding="utf-8") as f:
                    lines = f.readlines()

                if not lines:
                    continue

                # First line = title
                # title = lines[0].strip()
                title = lines[0].strip() if lines else path.stem

                # Remaining lines = body
                body = "".join(lines[1:]).strip()

                if not body:
                    continue

                docs.append(
                    Document(
                        page_content=body,
                        metadata={
                            "source": path.name,
                            "file_type": "txt",
                            "title": title,
                            "section": None,
                            "page": None,
                            "corpus": corpus_name,
                        },
                    )
                )

                print(f"Loaded: {path.name}")
                print(f"Total documents so far: {len(docs)}\n")

            # ---------------- PDF ----------------
            elif suffix == ".pdf":
                loader = PyPDFLoader(str(path))
                loaded = loader.load()

                for doc in loaded:
                    doc.metadata.update({
                        "source": path.name,
                        "file_type": "pdf",
                        "title": path.stem,
                        # keep page number from PyPDFLoader
                        "corpus": corpus_name,
                    })

                docs.extend(loaded)
                print(f"Loaded: {path.name}")
                print(f"Total documents so far: {len(docs)}\n")


            # ---------------- DOCX ----------------
            # note: convert .doc to .docx
            elif suffix == ".docx":
                loaded = load_docx_with_sections(path, corpus_name)
                docs.extend(loaded)
                print(f"Loaded: {path.name}")
                print(f"Total documents so far: {len(docs)}\n")

            # ---------------- DOC (skip) ----------------
            elif suffix == ".doc":
                print(f"Skipping legacy .doc file (convert to .docx): {path.name}")
                continue

            # ---------------- ODT ----------------
             
            elif suffix == ".odt":
                print(f"Converting ODT to DOCX: {path.name}")

#                converted_path = path.with_suffix(".docx")
                tmp_path=None
                try:
                    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                        tmp_path = Path(tmp.name)
                    
                    subprocess.run(
                        [
                            "pandoc",
                            "--from=odt",
                            "--to=docx",
                            str(path),
                            "--output",
                            str(tmp_path),
                        ],
                        check=True,
                        capture_output=True
                    )

                    loaded = load_docx_with_sections(tmp_path, corpus_name)
                    for doc in loaded:
                        doc.metadata["file_type"] = "odt"
                        doc.metadata["source"] = path.name  # preserve original name
               
                    docs.extend(loaded)

                    print(f"Loaded (via DOCX conversion): {path.name}")
                    print(f"Total documents so far: {len(docs)}\n")   
                         
                finally:
                    tmp_path.unlink(missing_ok=True)
                
            # ---------------- MD ----------------
            elif suffix == ".md":
                loaded = load_markdown_ast(path, corpus_name)
                docs.extend(loaded)
                print(f"Loaded: {path.name}")
                print(f"Total documents so far: {len(docs)}\n") 
 
            # --------------- JSON (Converation)---------------               
            elif suffix == ".json":
                loaded = load_json_conversation(path, corpus_name="memory")
                docs.extend(loaded)
                print(f"Loaded conversation: {path.name}")
                print(f"Total documents so far: {len(docs)}\n")              
                
            # ----------------- HTML --------------
            elif suffix in (".html", ".htm"):
                loaded = load_html_with_metadata(path, corpus_name)
                docs.extend(loaded)
                print(f"Loaded: {path.name}")
                print(f"Total documents so far: {len(docs)}\n")

                continue

            else:
                continue
                


#            print(f"Loaded: {path.name}")
#            print(f"Total documents so far: {len(docs)}")

        except Exception as e:
            print(f"Failed to load {path.name}: {e}")

    return docs
