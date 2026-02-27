#!/usr/bin/env python3
# research_assistant_v8.py
# look in religion directory

# v4 add Hybrid search
# v5 add session Memory
# v6 add conversation logger
# v7 add configuration file
# v8 add command line arguments
# v9 add json, md(markdown) document loaders, corpus configuration


import os
from pathlib import Path
from datetime import datetime

from langchain_core.documents import Document
from langchain_community.document_loaders import (
    TextLoader,
    UnstructuredODTLoader,
    PyPDFLoader,
#    BSHTMLLoader,
)
from langchain_text_splitters import RecursiveCharacterTextSplitter

from langchain_ollama import OllamaEmbeddings, ChatOllama
from langchain_community.vectorstores import FAISS
from langchain_community.retrievers import BM25Retriever

from langchain_core.prompts import ChatPromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.runnables import RunnableLambda
from langchain_core.chat_history import InMemoryChatMessageHistory
from langchain_core.runnables.history import RunnableWithMessageHistory

from docx import Document as DocxDocument
from bs4 import BeautifulSoup
from markdown_it import MarkdownIt

import json
import copy
import argparse

# from langchain_core.documents import Document

VERSION="0.0.9"

# =============================
# Configuration
# =============================

# PAPERS_DIR = Path("/home/snyder/ollama/religion/BofConcord")
# DB_DIR = Path("religion/BofC_faiss_index")
# EMBED_MODEL_NAME = "nomic-embed-text"
# CHAT_MODEL_NAME = "qwen3:4b-q4_K_M"

DEFAULT_CONFIG = {
    "paths": {
        "papers_dir": "./papers",
        "vectorstore_dir": "./faiss_index",
        "conversation_dir": "./conversations",
        "retrieval_log_dir": "./retrieval_logs"
    },
    "models": {
        "embedding_model": "nomic-embed-text",
        "chat_model": "qwen3:4b-q4_K_M",
        "temperature": 0.2
    },
    "retrieval": {
        "vector_k": 4,
        "bm25_k": 4
    },
    "chunking": {
        "chunk_size": 800,
        "chunk_overlap": 150
    },
    "logging": {
        "log_conversations": False,
        "log_retrieval_metadata": False
    },
    "session": {
        "default_session_id": "default"
    },
    "corpus":"research"
}

def load_config(path="config.json"):
    if not Path(path).exists():
        print("No config.json found. Using defaults.")
        return copy.deepcopy(DEFAULT_CONFIG)

    with open(path) as f:
        user_config = json.load(f)

    return deep_merge(copy.deepcopy(DEFAULT_CONFIG), user_config)
            
def deep_merge(default, override):
    for key, value in override.items():
        if (
            key in default
            and isinstance(default[key], dict)
            and isinstance(value, dict)
        ):
            default[key] = deep_merge(default[key], value)
        else:
            default[key] = value
    return default

def load_html_with_metadata(path, corpus_name):
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        soup = BeautifulSoup(f, "lxml")

    #title = soup.title.string.strip() if soup.title else path.name
    title = (
        soup.title.get_text(strip=True)
        if soup.title and soup.title.get_text(strip=True)
        else path.stem
    )

    docs = []

    for header in soup.find_all(["h1", "h2", "h3"]):
        section_title = header.get_text(strip=True)

        content = []
        for sibling in header.next_siblings:
            if hasattr(sibling, "name") and sibling.name in ["h1", "h2", "h3"]:

                break
            if hasattr(sibling, "get_text"):
                content.append(sibling.get_text(" ", strip=True))

        text = "\n".join(content).strip()

        if text:
            docs.append(
                Document(
                    page_content=text,
                    metadata={
                        "source": path.name,
                        "title": title,
                        "section": section_title,
                        "page": None,
                        "file_type": "html",
                        "corpus": corpus_name,
                    }
                )
            )

    return docs

    
# =============================
# Markdown(md) AST Loader
# =============================

def load_markdown_ast(path: Path, corpus_name):
    docs = []

    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        text = f.read()

    if not text.strip():
        return docs

    md = MarkdownIt()
    tokens = md.parse(text)

    title = path.stem
    current_headings = {1: None, 2: None, 3: None}
    buffer = []
    inside_code_block = False
    code_language = None

    def current_section():
        # Use deepest available heading
        for level in reversed([1, 2, 3]):
            if current_headings[level]:
                return current_headings[level]
        return "Introduction"

    def flush_section():
        if buffer:
            docs.append(
                Document(
                    page_content="\n".join(buffer).strip(),
                    metadata={
                        "source": path.name,
                        "file_type": "md",
                        "title": title,
                        "section": current_section(),
                        "page": None,
                        "contains_code": inside_code_block,
                        "code_language": code_language,
                        "corpus": corpus_name,
                    },
                )
            )

    i = 0
    while i < len(tokens):
        token = tokens[i]

        # ---------------- Headings ----------------
        if token.type == "heading_open":
            flush_section()
            buffer.clear()

            level = int(token.tag[1])  # h1, h2, h3
            heading_text = tokens[i + 1].content

            if level == 1:
                title = heading_text
                current_headings = {1: heading_text, 2: None, 3: None}
            else:
                current_headings[level] = heading_text

            i += 2  # skip inline + heading_close

        # ---------------- Paragraph / Inline ----------------
        elif token.type == "inline":
            buffer.append(token.content)

        # ---------------- Code Block ----------------
        elif token.type == "fence":
            flush_section()
            buffer.clear()

            inside_code_block = True
            code_language = token.info.strip() or None

            docs.append(
                Document(
                    page_content=token.content,
                    metadata={
                        "source": path.name,
                        "file_type": "md",
                        "title": title,
                        "section": current_section(),
                        "page": None,
                        "contains_code": True,
                        "code_language": code_language,
                        "corpus": corpus_name,
                    },
                )
            )

            inside_code_block = False
            code_language = None

        # ---------------- Lists ----------------
        elif token.type == "list_item_open":
            buffer.append("- ")

        i += 1

    flush_section()

    return docs   

# =============================
# Custom DOCX Loader (with sections + title)
# =============================

def load_docx_with_sections(path: Path, corpus_name):
    docs = []
    doc = DocxDocument(str(path))

    # Extract Word metadata title (fallback to filename)
    core_props = doc.core_properties
    title = core_props.title if core_props.title else path.stem

    current_section = "Introduction"

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style.name.startswith("Heading"):
            current_section = text
            continue

        docs.append(
            Document(
                page_content=text,
                metadata={
                    "source": path.name,
                    "file_type": "docx",
                    "title": title,
                    "section": current_section,
                    "corpus": corpus_name,
                },
            )
        )

    return docs

# =============================
# JSON Conversation Loader
# =============================

import json

def load_json_conversation(path: Path):
    docs = []

    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)

    session_id = data.get("session_id", path.stem)
    messages = data.get("messages", [])

    i = 0
    while i < len(messages):
        msg = messages[i]

        # Pair user + assistant if possible
        if msg.get("role") == "user":
            user_content = msg.get("content", "")
            assistant_content = ""

            if i + 1 < len(messages) and messages[i + 1].get("role") == "assistant":
                assistant_content = messages[i + 1].get("content", "")
                i += 1  # skip assistant

            combined = (
                f"User:\n{user_content}\n\n"
                f"Assistant:\n{assistant_content}"
            )

            docs.append(
                Document(
                    page_content=combined.strip(),
                    metadata={
                        "source": path.name,
                        "file_type": "json_conversation",
                        "session_id": session_id,
                        "page": None,
                        "corpus": "memory",  # explicitly different
                    },
                )
            )

        i += 1

    return docs
    

# =============================
# Step 1: Load documents
# =============================

def load_documents(folder: Path, corpus_name):
    docs = []

    for path in folder.glob("**/*"):
        suffix = path.suffix.lower()

        try:

            # ---------------- TXT ----------------
            if suffix == ".txt":

                with open(path, "r", encoding="utf-8") as f:
                    lines = f.readlines()

                if not lines:
                    continue

                # First line = title
                # title = lines[0].strip()
                title = lines[0].strip() if lines else path.stem

                # Remaining lines = body
                body = "".join(lines[1:]).strip()

                if not body:
                    continue

                docs.append(
                    Document(
                        page_content=body,
                        metadata={
                            "source": path.name,
                            "file_type": "txt",
                            "title": title,
                            "section": None,
                            "page": None,
                            "corpus": corpus_name,
                        },
                    )
                )

#                print(f"Loaded: {path.name}")
#                print(f"Total documents so far: {len(docs)}\n")

            # ---------------- PDF ----------------
            elif suffix == ".pdf":
                loader = PyPDFLoader(str(path))
                loaded = loader.load()

                for doc in loaded:
                    doc.metadata.update({
                        "source": path.name,
                        "file_type": "pdf",
                        "title": path.stem,
                        # keep page number from PyPDFLoader
                        "corpus": corpus_name,
                    })

                docs.extend(loaded)
                print(f"Loaded: {path.name}")
                print(f"Total documents so far: {len(docs)}\n")


            # ---------------- DOCX ----------------
            # note: convert .doc to .docx
            elif suffix == ".docx":
                loaded = load_docx_with_sections(path, corpus_name)
                docs.extend(loaded)
                print(f"Loaded: {path.name}")
                print(f"Total documents so far: {len(docs)}\n")

            # ---------------- DOC (skip) ----------------
            elif suffix == ".doc":
                print(f"Skipping legacy .doc file (convert to .docx): {path.name}")
                continue

            # ---------------- ODT ----------------
            elif suffix == ".odt":
                loader = UnstructuredODTLoader(str(path))
                loaded = loader.load()

                for doc in loaded:
                    doc.metadata.update({
                        "source": path.name,
                        "file_type": "odt",
                        "title": path.stem,
                        "corpus": corpus_name,
                    })

                docs.extend(loaded)
                print(f"Loaded: {path.name}")

            # ---------------- MD ----------------
            elif suffix == ".md":
                loaded = load_markdown_ast(path, corpus_name)
                docs.extend(loaded)
                print(f"Loaded: {path.name}")
                print(f"Total documents so far: {len(docs)}\n") 
 
            # --------------- JSON (Converation)---------------               
            elif suffix == ".json":
                loaded = load_json_conversation(path)
                docs.extend(loaded)
                print(f"Loaded conversation: {path.name}")
                print(f"Total documents so far: {len(docs)}\n")              
                
            # ----------------- HTML --------------
            elif suffix in (".html", ".htm"):
                loaded = load_html_with_metadata(path, corpus_name)
                docs.extend(loaded)
                print(f"Loaded: {path.name}")
                print(f"Total documents so far: {len(docs)}\n")

                continue

            else:
                continue
                


            print(f"Loaded: {path.name}")
            print(f"Total documents so far: {len(docs)}")

        except Exception as e:
            print(f"Failed to load {path.name}: {e}")

    return docs


# =============================
# Step 2: Chunk documents
# =============================

def chunk_documents(docs, config):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=config["chunking"]["chunk_size"],
        chunk_overlap=config["chunking"]["chunk_overlap"],
    )

    chunks = splitter.split_documents(docs)

    # Ensure metadata persists
    for chunk in chunks:
        chunk.metadata.setdefault("section", None)
        chunk.metadata.setdefault("page", None)

    return chunks


# =============================
# Step 3: Create or load FAISS
# =============================


    
def build_vectorstore(chunks, config):
    embeddings = OllamaEmbeddings(
        model=config["models"]["embedding_model"]
    )

    db_dir = Path(config["paths"]["vectorstore_dir"])

    vectorstore = FAISS.from_documents(chunks, embeddings)
    vectorstore.save_local(str(db_dir))
    return vectorstore
    
# Build file manifest to dectect changes

def build_file_manifest(folder: Path):
    manifest = {}

    for path in folder.glob("**/*"):
        if path.is_file():
            manifest[str(path)] = {
                "mtime": path.stat().st_mtime,
                "size": path.stat().st_size,
            }

    return manifest
    
# def needs_rebuild(config):
#
#    db_dir = Path(config["paths"]["vectorstore_dir"])
#    papers_dir = Path(config["paths"]["papers_dir"])
#    manifest_path = db_dir / "manifest.json"
#
#    if not manifest_path.exists():
#        return True
#
#    with open(manifest_path, "r") as f:
#        old_manifest = json.load(f)
#
#    current_manifest = build_file_manifest(papers_dir)
#    return old_manifest != current_manifest

def needs_rebuild(config):

    db_dir = Path(config["paths"]["vectorstore_dir"])
    papers_dir = Path(config["paths"]["papers_dir"])
    manifest_path = db_dir / "manifest.json"

    # No manifest → rebuild
    if not manifest_path.exists():
        return True

    try:
        with open(manifest_path, "r") as f:
            saved_manifest = json.load(f)
    except Exception:
        return True  # corrupted manifest → rebuild

    # Check corpus change
    current_corpus = config.get("corpus", "research")
    if saved_manifest.get("corpus") != current_corpus:
        return True

    # Check file changes
    current_files = build_file_manifest(papers_dir)
    if saved_manifest.get("files") != current_files:
        return True

    return False

# =============================
# Step 4: Build LCEL RAG chain
# =============================

def format_docs(docs):
    formatted = []
    for d in docs:
        meta = d.metadata

        citation_parts = []

        if meta.get("title"):
            citation_parts.append(meta["title"])

        if meta.get("source"):
            citation_parts.append(meta["source"])

        if meta.get("section"):
            citation_parts.append(f"section: {meta['section']}")

        if meta.get("page"):
            citation_parts.append(f"page: {meta['page']}")

        citation = "  ".join(citation_parts)

        formatted.append(f"""
[{citation}]
{d.page_content}
""")

    return "\n\n".join(formatted)


def create_rag_chain(vectorstore, chunks, config):

    llm = ChatOllama(
        model=config["models"]["chat_model"],
        temperature=config["models"]["temperature"]
    )

    vector_retriever = vectorstore.as_retriever(
        search_kwargs={"k": config["retrieval"]["vector_k"]}
    )

    bm25_retriever = BM25Retriever.from_documents(chunks)
    bm25_retriever.k = config["retrieval"]["bm25_k"]

    def hybrid_retriever_fn(query):
        bm25_docs = bm25_retriever.invoke(query)
        vector_docs = vector_retriever.invoke(query)

        all_docs = bm25_docs + vector_docs
        unique_docs = {doc.page_content: doc for doc in all_docs}.values()
        return list(unique_docs)

    def retrieval_with_logging(x):
        query = x["question"]
        session_id = x.get("session_id", "default")

        docs = hybrid_retriever_fn(query)
        log_retrieval(config, session_id, query, docs)
        return docs

    hybrid_runnable = RunnableLambda(retrieval_with_logging)
    formatter_runnable = RunnableLambda(format_docs)

    prompt = ChatPromptTemplate.from_template("""
You are a research assistant.

Use ONLY the provided context to answer the question.

When you use information from the context:
- Cite the source in parentheses.
- Include title, source filename, section (if available), and page (if available).

Format citation like:
(Title  source, section: X, page: Y)

If section or page is missing, omit it.

At the end of the answer, include a "Sources:" list
with one entry per unique document used.

Conversation History:
{history}

Context:
{context}

Question:
{question}
""")

    rag_chain = (
        {
            "context": hybrid_runnable | formatter_runnable,
            "question": RunnableLambda(lambda x: x["question"]),
            "history": RunnableLambda(lambda x: x.get("history", "")),
        }
        | prompt
        | llm
    )

    store = {}

    def get_session_history(session_id):
        if session_id not in store:
            store[session_id] = InMemoryChatMessageHistory()
        return store[session_id]

    return RunnableWithMessageHistory(
        rag_chain,
        get_session_history,
        input_messages_key="question",
        history_messages_key="history",
    )
    
 
# =============================
# Step 5: Interactive session
# =============================


def interactive_session(rag_chain, config):

    print("\nResearch Assistant Ready!")
    print("Type 'exit' or 'quit' to stop.\n")

    default_session = config["session"]["default_session_id"]
    session_id = input("Session name: ").strip() or default_session
    while True:
        query = input("Ask a question about the documents: ").strip()
        if query.lower() in ("exit", "quit"):
            break

        response = rag_chain.invoke(
            {"question": query},
            config={"configurable": {"session_id": session_id}}
        )

        print("\n", response.content, "\n")

        save_conversation(
            config,
            session_id,
            query,
            response.content
        )
                
# conversation Logger
        
# CONVO_DIR = Path("conversations")
# CONVO_DIR.mkdir(exist_ok=True)

def save_conversation(config, session_id, question, answer):
    if not config["logging"]["log_conversations"]:
        return

    convo_dir = Path(config["paths"]["conversation_dir"])
    convo_dir.mkdir(exist_ok=True)

    filepath = convo_dir / f"{session_id}.json"

    if filepath.exists():
        try:
            with open(filepath, "r") as f:
                data = json.load(f)
        except json.JSONDecodeError:
            data = []
    else:
        data = []

    data.append({
        "timestamp": datetime.now().isoformat(),
        "model": config["models"]["chat_model"],
        "question": question,
        "answer": answer,
    })

    with open(filepath, "w") as f:
        json.dump(data, f, indent=2)



# log retrieval metadata
# RETRIEVAL_LOG = Path("retrieval_logs")
# RETRIEVAL_LOG.mkdir(exist_ok=True)

def log_retrieval(config, session_id, query, docs):
    if not config["logging"]["log_retrieval_metadata"]:
        return

    retrieval_dir = Path(config["paths"]["retrieval_log_dir"])
    retrieval_dir.mkdir(exist_ok=True)

    filepath = retrieval_dir / f"{session_id}_retrieval.json"

    if filepath.exists():
        with open(filepath, "r") as f:
            data = json.load(f)
    else:
        data = []

    data.append({
        "timestamp": datetime.now().isoformat(),
        "query": query,
        "documents": [
            {
                "title": d.metadata.get("title"),
                "source": d.metadata.get("source"),
                "section": d.metadata.get("section"),
                "page": d.metadata.get("page"),
            }
            for d in docs
        ],
    })

    with open(filepath, "w") as f:
        json.dump(data, f, indent=2)

   

# =============================
# Command Line Arguments
# =============================

def parse_args():
    parser = argparse.ArgumentParser(
        prog="research_assistant.py",
        description="Local Hybrid RAG Research Assistant using LangChain + Ollama.",
        epilog="CLI overrides config.json. Config overrides defaults.",
        formatter_class=argparse.ArgumentDefaultsHelpFormatter,
    )

    parser.add_argument(
        "-c", "--config",
        default="config.json",
        help="Path to configuration file"
    )

    parser.add_argument(
        "-m", "--model",
        help="Override LLM chat model (e.g. qwen3:4b-q4_K_M)"
    )

    parser.add_argument(
        "-d", "--documents",
        help="Override papers directory"
    ) 
    
    # ✅ Version flag
    parser.add_argument(
        "--version",
        action="version",
        version=f"%(prog)s {VERSION}",
    )

    
    

    return parser.parse_args() 

# =============================
# Main
# =============================

if __name__ == "__main__":

    args = parse_args()

    # Load config file (or defaults)
    config = load_config(args.config)

    # =============================
    # CLI Overrides (highest priority)
    # =============================

    if args.model:
        config["models"]["chat_model"] = args.model

    if args.documents:
        config["paths"]["papers_dir"] = args.documents

    # Optional: show effective config
    print("\nEffective configuration:")
    print(json.dumps(config, indent=2))

    papers_dir = Path(config["paths"]["papers_dir"])
    db_dir = Path(config["paths"]["vectorstore_dir"])
    
    corpus_name = config["corpus"]

    if needs_rebuild(config):
        print("Corpus changed  rebuilding index.")

        documents = load_documents(papers_dir, corpus_name)
        chunks = chunk_documents(documents, config)

        # delete old index if exists
        if db_dir.exists():
            import shutil
            shutil.rmtree(db_dir)


        # rebuid manifest
        print("rebuilding vectorstore ...")
        db_dir.mkdir(exist_ok=True)
        vectorstore = build_vectorstore(chunks, config)
        manifest = {
            "corpus": corpus_name,
            "files": build_file_manifest(papers_dir)
        }

        db_dir.mkdir(exist_ok=True)         
        with open(db_dir / "manifest.json", "w") as f:
            json.dump(manifest, f, indent=2)

    else:
        print("Corpus unchanged  loading existing index.")
        embeddings = OllamaEmbeddings(
            model=config["models"]["embedding_model"]
        )
        vectorstore = FAISS.load_local(
            str(db_dir),
            embeddings,
            allow_dangerous_deserialization=True
        )

        # 🔹 Still reload docs for BM25
        documents = load_documents(papers_dir, corpus_name)
        chunks = chunk_documents(documents,config)
 
    rag_chain = create_rag_chain(vectorstore, chunks, config)
   
    interactive_session(rag_chain,config)
